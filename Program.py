import os
from PyPDF2 import PdfReader
from docx import Document
from langchain_text_splitters import TokenTextSplitter

def extract_text_from_pdf(file_path):
    text = ''
    with open(file_path,'rb') as pdf_file:
        pdf_reader = PdfReader(pdf_file)
        page_nums = len(pdf_reader.pages)

        for page_num in range(page_nums):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() 
    return text


def extract_text_from_docx(file_path):
    document = Document(file_path)
    text = ''
    for paragraph in document.paragraphs:
        if len(paragraph.text)>0:
            text += paragraph.text 

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text 
    return text

  
def process_folder(folder_path):
    text_variable = []
    for root, dirs, files in os.walk(folder_path):
        for file_name in files:
            file_path = os.path.join(root, file_name)
            text = ''
            if file_name.lower().endswith('.docx'):
                text = extract_text_from_docx(file_path)
            elif file_name.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)  
            else:
                continue
            
            if not isinstance(text, str):
                continue
            
            token_splitter = TokenTextSplitter(chunk_size = 500, chunk_overlap = 50)
            tokens = token_splitter.split_text(text)
            
            metadata = {
                'filename': file_name,
                'folder_path': root,
                'subfolders': os.path.relpath(root, folder_path).split(os.sep)[1:],
                'tokens': tokens
            }
            text_variable.append(metadata)
    return text_variable


folder_path = r'C:\Users\MiquelMuletAlarcón\Desktop\Pyhton\ha\hackbcnai2024'
processed_data = process_folder(folder_path)

