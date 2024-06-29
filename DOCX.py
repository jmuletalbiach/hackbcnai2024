from docx import Document

def extreure_text_docx(nom_fitxer):
    document = Document(nom_fitxer)
    text = []

    for paragraph in document.paragraphs:
        if (len(paragraph.text)>0):
            text.append(paragraph.text)

    for tables in document.tables:
        for fila in tables.rows:
            for cell in fila.cells:
                text.append(cell.text)

    text_complet = "\n".join(text)
    return text_complet

nom_fitxer = 'Prova DOCX.docx'
text = extreure_text_docx(nom_fitxer)
print(text)





