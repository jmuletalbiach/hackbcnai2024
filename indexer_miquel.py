import os
from PyPDF2 import PdfReader
from docx import Document
from os import listdir
from os.path import isfile, join,isdir
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_qdrant import Qdrant
import sys
from langchain_text_splitters import TokenTextSplitter
from pptx import Presentation
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams
import docx
import environment_var

def get_input_streams(dir):
    stream_list = []
    for f in listdir(dir):
        if isfile(join(dir,f)):
            stream_list.append(join(dir,f))
        elif isdir(join(dir,f)):
            stream_list= stream_list + get_input_streams(join(dir,f))
    return stream_list




def main_indexer(folder_path):
    nom_model = "sentence-transformers/msmarco-bert-base-dot-v5"
    arguments_model = {'device': 'cpu'}
    arguments_encoder = {'normalize_embeddings': True}
    hf = HuggingFaceEmbeddings(
        model_name=nom_model,
        model_kwargs=arguments_model,
        encode_kwargs=arguments_encoder
    )

    if environment_var.qdrant_storage =="cloud":
        client = QdrantClient(
        url=environment_var.qdrant_cloud_url, 
        api_key=environment_var.qdrant_cloud_key,
        )
        print ("qdrant cloud service instance")
        print(client.get_collections())
    elif environment_var.qdrant_storage =="local":
        client = QdrantClient("localhost", port=6333)
        print ("qdrant local instance")
        print(client.get_collections())
    

    collection_name = "SocialEducation"
    if environment_var.qdrant_clean == "yes":
        if client.collection_exists(collection_name):
            client.delete_collection(collection_name)
        client.create_collection(collection_name,vectors_config=VectorParams(size=768, distance=Distance.DOT))
    qdrant = Qdrant(client, collection_name, hf)
    print(client.get_collections())

    print("Retrieving stream sources")
    text_variable = []
    for root, dirs, files in os.walk(folder_path):
        for file_name in files:
            file_path = os.path.join(root, file_name)
            text = ''
            if file_name.lower().endswith('.docx'):
                text = extract_text_from_docx(file_path)
            elif file_name.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)  
            else:
                continue
            
            if not isinstance(text, str):
                continue
            
            subfolders = os.path.relpath(root, folder_path).split(os.sep)[0:]
            if len(subfolders)<2:
                subfolders.extend([' '])
            
            try:
                token_splitter = TokenTextSplitter(chunk_size=500, chunk_overlap=50)
                tokens = token_splitter.split_text(text)
            except Exception as e:
                continue
            
            metadata = []
            #adding path metadata
            for i in range(0,len(tokens)):
                metadata.append({'filename': file_name})
                metadata.append({'subfolders': subfolders})
            qdrant.add_texts(tokens,metadatas=metadata)
            print(file_name)
    
    print("Incremental stream processing processed!")



def extract_text_from_pdf(file_path):
    text = ''
    with open(file_path,'rb') as pdf_file:
        pdf_reader = PdfReader(pdf_file)
        page_nums = len(pdf_reader.pages)

        for page_num in range(page_nums):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() 
    return text


def extract_text_from_docx(file_path):
    document = Document(file_path)
    text = ''
    for paragraph in document.paragraphs:
        if len(paragraph.text)>0:
            text += paragraph.text 

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text 
    return text




if __name__ == "__main__":
    arguments = sys.argv
    if len(arguments)>1:
        main_indexer(arguments[1])
    else:
        print("Indexing is not possible. Please review input streams")